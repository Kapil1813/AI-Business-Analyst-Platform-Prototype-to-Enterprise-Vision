from docx import Document

def generate_brd(objective, asset_class, description, ai_output):

    doc = Document()

    doc.add_heading("Business Requirements Document", level=1)

    doc.add_heading("Business Objective", level=2)
    doc.add_paragraph(objective)

    doc.add_heading("Asset Class", level=2)
    doc.add_paragraph(asset_class)

    doc.add_heading("Request Description", level=2)
    doc.add_paragraph(description)

    doc.add_heading("AI Generated Requirements", level=2)
    doc.add_paragraph(ai_output)

    filename = "BRD_output.docx"
    doc.save(filename)

    return filename