# ai_engine.py
import openai
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import altair as alt
import streamlit as st

# --- Analyze Request via OpenAI ---
def analyze_request(request_text, api_key):
    """
    Calls OpenAI GPT to generate a structured BRD-ready requirements dictionary.
    """
    openai.api_key = api_key

    prompt = f"""
You are a senior business analyst. A user provides the following request:

{request_text}

Generate structured requirements text suitable for a BRD with the following sections:
1. Functional Requirements
2. Non-Functional Requirements
3. User Stories
4. Architecture Suggestions

Use clear numbering, bullets, and headings. Make it professional and concise.
"""
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful business analyst."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error generating requirements: {e}"

# --- Generate BRD Word Document ---
def generate_brd(business_objective, asset_class, request_description, ai_output):
    """
    Generates a structured BRD Word document with proper headings and bullets.
    Returns (bytes, filename) for Streamlit download.
    """
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading("Business Requirements Document", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Business Objective ---
    doc.add_heading("Business Objective", level=1)
    doc.add_paragraph(business_objective)

    # --- Asset Class ---
    doc.add_heading("Asset Class", level=1)
    doc.add_paragraph(asset_class)

    # --- Request Description ---
    doc.add_heading("Request Description", level=1)
    doc.add_paragraph(request_description)

    # --- AI Generated Requirements ---
    doc.add_heading("AI Generated Requirements", level=1)

    # Parse AI output into lines and create headings/bullets
    for line in ai_output.split("\n"):
        line = line.strip()
        if line.startswith("###") or line.startswith("1.") or line.startswith("2.") or line.startswith("3.") or line.startswith("4."):
            doc.add_heading(line.replace("#","").strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            doc.add_paragraph(line)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read(), "Business_Requirements_Document.docx"

# --- Dashboard placeholder ---
def display_dashboard():
    st.subheader("Requirements Dashboard")
    data = pd.DataFrame({
        "Type": ["Functional", "Non-Functional", "User Stories", "Architecture Suggestions"],
        "Count": [5, 5, 5, 5]  # placeholder counts
    })
    chart = alt.Chart(data).mark_bar().encode(
        x="Type:N",
        y="Count:Q",
        color="Type:N"
    )
    st.altair_chart(chart, use_container_width=True)